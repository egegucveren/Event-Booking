from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("/Users/egegucveren/Desktop/Event-Booking/PulsePass_Project_Documentation.docx")

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(142, 108, 45)
INK = RGBColor(33, 37, 41)
MUTED = RGBColor(99, 115, 129)
LIGHT_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "EAF1F8"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")


def apply_font(run, *, name: str = "Calibri", size: int = 11, color: RGBColor = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, *, before: int = 0, after: int = 6, line: float = 1.1, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text_paragraph(doc: Document, text: str, *, style: str = "Normal", before: int = 0, after: int = 6, line: float = 1.1, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = doc.add_paragraph(style=style)
    format_paragraph(paragraph, before=before, after=after, line=line, align=align)
    run = paragraph.add_run(text)
    style_font = {
        "Title": ("Calibri", 26, NAVY, True),
        "Subtitle": ("Calibri", 13, MUTED, False),
        "Heading 1": ("Calibri", 16, BLUE, True),
        "Heading 2": ("Calibri", 13, BLUE, True),
        "Heading 3": ("Calibri", 12, DARK_BLUE, True),
    }
    if style in style_font:
        name, size, color, bold = style_font[style]
        apply_font(run, name=name, size=size, color=color, bold=bold)
    else:
        apply_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        format_paragraph(paragraph, before=0, after=4, line=1.15)
        run = paragraph.add_run(item)
        apply_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        format_paragraph(paragraph, before=0, after=4, line=1.15)
        run = paragraph.add_run(item)
        apply_font(run)


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_layout_fixed(table)

    for label, value in rows:
        row = table.add_row()
        left, right = row.cells
        set_cell_width(left, 1.7)
        set_cell_width(right, 4.8)
        set_cell_shading(left, LIGHT_FILL)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        format_paragraph(lp, after=2, line=1.0)
        format_paragraph(rp, after=2, line=1.0)
        lrun = lp.add_run(label)
        rrun = rp.add_run(value)
        apply_font(lrun, bold=True, color=DARK_BLUE)
        apply_font(rrun)


def add_two_column_feature_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    add_text_paragraph(doc, title, style="Heading 2", before=12, after=8)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_layout_fixed(table)

    header = table.rows[0].cells
    headers = ["Area", "Details"]
    widths = [2.2, 4.3]
    for cell, text, width in zip(header, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_BLUE_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        format_paragraph(p, after=2, line=1.0)
        run = p.add_run(text)
        apply_font(run, bold=True, color=DARK_BLUE)

    for left_text, right_text in rows:
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = cells[0].paragraphs[0]
        rp = cells[1].paragraphs[0]
        format_paragraph(lp, after=2, line=1.05)
        format_paragraph(rp, after=2, line=1.05)
        apply_font(lp.add_run(left_text), bold=True)
        apply_font(rp.add_run(right_text))


def add_header_and_footer(section, header_label: str) -> None:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    format_paragraph(hp, before=0, after=0, line=1.0)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(header_label)
    apply_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    format_paragraph(fp, before=0, after=0, line=1.0)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_run = fp.add_run("PulsePass Project Documentation")
    apply_font(page_run, size=9, color=MUTED)


def create_document() -> None:
    doc = Document()
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color, before, after in [
        ("Title", 26, NAVY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True if style_name != "Subtitle" else False
        style.font.color.rgb = color
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)

    section = doc.sections[0]
    add_header_and_footer(section, "Software Web Development Project Report")

    cover_kicker = doc.add_paragraph()
    format_paragraph(cover_kicker, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_font(cover_kicker.add_run("PROJECT DOCUMENTATION"), size=11, color=GOLD, bold=True)

    add_text_paragraph(
        doc,
        "PulsePass",
        style="Title",
        after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text_paragraph(
        doc,
        "Booking and Event Management Web Application",
        style="Subtitle",
        after=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    tagline = doc.add_paragraph()
    format_paragraph(tagline, before=0, after=22, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_font(
        tagline.add_run(
            "A full-stack Next.js and MySQL platform for organisers, attendees, and administrators."
        ),
        size=11,
        color=INK,
        italic=True,
    )

    add_metadata_table(
        doc,
        [
            ("Document Type", "Software Web Development project documentation"),
            ("Application Name", "PulsePass"),
            ("Technology Stack", "Next.js 15, React 19, TypeScript, MySQL, Zod"),
            ("Prepared On", date(2026, 5, 9).strftime("%d %B %Y")),
            ("Primary Purpose", "Demonstrate project scope, architecture, features, and setup"),
        ],
    )

    spacer = doc.add_paragraph()
    format_paragraph(spacer, before=0, after=18)
    cover_note = doc.add_paragraph()
    format_paragraph(cover_note, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_font(
        cover_note.add_run(
            "This report summarises the system design, user workflows, database model, setup process, and assignment-oriented feature coverage of the PulsePass platform."
        ),
        size=10,
        color=MUTED,
    )

    doc.add_page_break()

    add_text_paragraph(doc, "Executive Summary", style="Heading 1")
    add_text_paragraph(
        doc,
        "PulsePass is a responsive event booking platform designed to support three core actors: organisers who publish and manage events, attendees who discover and reserve seats, and administrators who supervise users, platform activity, and support tickets. The application combines server-rendered and interactive user experiences with a MySQL-backed data model and validation-first server actions.",
        after=8,
    )
    add_text_paragraph(
        doc,
        "The project demonstrates practical full-stack development concerns including authentication, role-based access control, event lifecycle management, booking concurrency protection, database transactions, responsive interface design, and maintainable modular code organisation.",
        after=8,
    )

    add_text_paragraph(doc, "Project Objectives", style="Heading 2")
    add_bullets(
        doc,
        [
            "Provide a public storefront for browsing upcoming events with search and filter tools.",
            "Allow organisers to create, edit, monitor, and remove event listings from a dedicated dashboard.",
            "Allow attendees to register, log in, manage bookings, and maintain an active e-ticket card.",
            "Allow admins to manage users, delete platform events, and resolve support tickets.",
            "Demonstrate secure data handling, strong validation, and clean project structure for assessment purposes.",
        ],
    )

    add_text_paragraph(doc, "System Overview", style="Heading 1")
    add_text_paragraph(
        doc,
        "The application is built with the Next.js App Router and uses server components for route-level data composition, React client components for interactive forms and filters, and server actions to process authenticated mutations. Data is stored in MySQL tables for users, sessions, events, bookings, e-ticket cards, and contact tickets.",
        after=8,
    )
    add_two_column_feature_table(
        doc,
        "Architecture Summary",
        [
            ("Frontend", "Next.js App Router pages, reusable UI components, responsive CSS, client-side filters, and action-driven forms."),
            ("Backend", "Server actions for authentication, events, bookings, admin actions, contact submissions, and e-ticket operations."),
            ("Database", "MySQL schema with relational constraints, indexes, and transactional booking logic."),
            ("Security", "Scrypt password hashing, HTTP-only session cookies, SHA-256 token hashing, and role-based route protection."),
            ("Validation", "Zod schemas validate registration, login, events, bookings, user changes, and support messages."),
        ],
    )

    add_text_paragraph(doc, "User Roles and Core Workflows", style="Heading 1")
    add_two_column_feature_table(
        doc,
        "Role Responsibilities",
        [
            ("Organiser", "Creates events, edits event details, monitors bookings, and removes their own listings."),
            ("Attendee", "Searches the marketplace, books seats, manages reservations, and maintains an e-ticket card."),
            ("Admin", "Manages user roles, deletes platform events when necessary, and resolves open contact tickets."),
        ],
    )
    add_text_paragraph(doc, "Typical End-to-End Flow", style="Heading 2")
    add_numbered(
        doc,
        [
            "A visitor browses public event listings through the marketplace page.",
            "A new user creates an organiser or attendee account and signs in.",
            "Organisers publish event listings with structured details such as city, venue, category, pricing, and capacity.",
            "Attendees open an event detail page, confirm a valid e-ticket card, and reserve up to six seats in one booking.",
            "Admins monitor platform activity from the control dashboard and can resolve tickets or remove problematic data.",
        ],
    )

    add_text_paragraph(doc, "Key Features", style="Heading 1")
    add_two_column_feature_table(
        doc,
        "Feature Coverage",
        [
            ("Authentication", "Registration, login, logout, password hashing, and session creation through server actions."),
            ("Role-Based Access", "Route protection redirects users to the correct dashboard based on their assigned role."),
            ("Event Management", "Create, update, cancel, preview, and delete event listings from organiser and admin workflows."),
            ("Booking Management", "Attendees can create and cancel bookings with duplicate booking and capacity checks."),
            ("Support Workflow", "Contact form submissions are stored in the database and can be resolved by administrators."),
            ("Marketplace UX", "Homepage includes live search plus city and category filtering for event discovery."),
            ("Responsive UI", "Layouts adapt across dashboards, cards, tables, and public-facing pages."),
        ],
    )

    add_text_paragraph(doc, "Database Design", style="Heading 1")
    add_text_paragraph(
        doc,
        "The relational data model supports the full booking lifecycle. Users can own events, attend bookings, hold sessions, and receive e-ticket cards. Events and bookings are linked with foreign keys, while contact tickets provide a lightweight support workflow for admin review.",
        after=8,
    )
    add_two_column_feature_table(
        doc,
        "Main Tables",
        [
            ("users", "Stores identity, email, password hash, role, ownership flag, and account creation date."),
            ("sessions", "Stores hashed session tokens, expiry timestamps, and links sessions to users."),
            ("events", "Stores organiser ownership, event metadata, timing, pricing, capacity, descriptions, and status."),
            ("e_ticket_cards", "Stores attendee-linked card status, issue date, and expiry date."),
            ("bookings", "Stores booking code, event link, attendee link, seats, totals, and booking state."),
            ("contact_tickets", "Stores sender details, message content, status, and submission timestamp."),
        ],
    )

    add_text_paragraph(doc, "Validation, Security, and Reliability", style="Heading 1")
    add_bullets(
        doc,
        [
            "Passwords are hashed with scrypt and session tokens are stored in the database as SHA-256 hashes.",
            "Session cookies are HTTP-only and use expiry-based server validation.",
            "Role checks are enforced before protected dashboard actions are executed.",
            "Zod schemas enforce field-level validation for all major forms and action payloads.",
            "Booking creation uses a transaction with SELECT FOR UPDATE to reduce race conditions around seat allocation.",
            "Confirmed bookings are cancelled automatically when an organiser or admin deletes an event.",
        ],
    )

    add_text_paragraph(doc, "Project Structure", style="Heading 1")
    add_two_column_feature_table(
        doc,
        "Repository Layout",
        [
            ("app/", "Route files for public pages, dashboards, and event detail pages."),
            ("actions/", "Server actions for authentication, events, bookings, e-ticket handling, admin tasks, and contact logic."),
            ("components/", "Reusable layout, form, and UI building blocks."),
            ("lib/", "Authentication helpers, database utilities, formatting helpers, validation schemas, and typed query helpers."),
            ("database/", "Schema and migration SQL for the application data model."),
            ("scripts/", "Database seed, migration scripts, and supporting automation."),
        ],
    )

    add_text_paragraph(doc, "Setup and Local Usage", style="Heading 1")
    add_numbered(
        doc,
        [
            "Install project dependencies with npm install.",
            "Copy .env.local.example to .env.local and provide the local MySQL configuration.",
            "Start the MySQL service on the development machine.",
            "Seed the database with npm run seed.",
            "Start the local development server with npm run dev.",
            "Use npm run typecheck to generate route types and verify TypeScript correctness.",
        ],
    )

    add_text_paragraph(doc, "Demo Accounts", style="Heading 2")
    add_two_column_feature_table(
        doc,
        "Login Reference",
        [
            ("admin@pulsepass.local", "Admin owner account"),
            ("organiser@pulsepass.local", "Organiser account"),
            ("attendee@pulsepass.local", "Attendee account"),
            ("Shared Password", "Passo123!"),
        ],
    )

    add_text_paragraph(doc, "Assignment-Oriented Coverage", style="Heading 1")
    add_two_column_feature_table(
        doc,
        "Requirement Mapping",
        [
            ("User registration and login", "Implemented through dedicated register and login flows backed by server-side validation."),
            ("Role-based access control", "Protected dashboards and actions route users according to admin, organiser, or attendee roles."),
            ("Event CRUD", "Organisers can create, edit, and delete events; admins can remove any event."),
            ("Booking management", "Attendees can create and cancel bookings while capacity is enforced transactionally."),
            ("User administration", "Admins can update roles and delete user accounts subject to ownership rules."),
            ("Contact ticket workflow", "Public contact submissions are stored and can be resolved from the admin dashboard."),
            ("Server-side validation", "Zod schemas cover forms for auth, events, bookings, admin actions, and contact tickets."),
            ("Session management", "Session cookies and hashed tokens are used to maintain authenticated access with 7-day expiry."),
            ("Modular structure", "The repo separates pages, actions, components, data utilities, and schema logic."),
            ("Responsive design", "Grid and card-based interfaces adapt across the public site and dashboards."),
        ],
    )

    add_text_paragraph(doc, "Contributions", style="Heading 1")
    add_text_paragraph(
        doc,
        "In this section, I summarise how we divided the work across the project. I based this breakdown on the recorded GitHub and repository commit history together with the contributor name clarifications we applied in the documentation. The percentages are approximate and are intended to reflect each person's main area of contribution as fairly as possible.",
        after=8,
    )
    add_two_column_feature_table(
        doc,
        "Division of Labour",
        [
            ("Ege Gucveren (approx. 35%)", "I led the project foundation and core architecture. My recorded GitHub work includes creating the initial full-stack baseline, setting up the main Next.js structure, shared styling system, authentication utilities, organiser and attendee dashboards, event CRUD foundation, reusable UI components, database schema, validation setup, seed workflow, and later refinements to homepage queries, event visuals, media layout, configuration, and general cleanup."),
            ("Muhammed Furkan Erdem / Kayky / mferdem7 (approx. 24%)", "Muhammed Furkan Erdem focused mainly on attendee-facing event discovery and booking experience improvements. His recorded repository work includes custom search result cards, event detail layout improvements, smart booking calls to action, the yearly e-ticket card flow, QR ticket support, and explanatory comments for the search and e-ticket logic. The commit history shows this work under both the Kayky and mferdem7 identities."),
            ("Eyupcan Dinsever / Eyupdzhan Dinsever / Heaven1402 (approx. 27%)", "Based on our clarification, I treated these commit identities as the same contributor. In the combined repository history, this contribution covers the contact workflow, admin-facing ticket handling, authentication reliability, booking race-condition fixes, owner-admin restrictions, login page demo account visibility, revalidation updates, README changes, and general bug-fix passes affecting admin, events, database access, and query behaviour."),
            ("Berykay Yalmaz (recorded in repository as Efe Girgin) (approx. 14%)", "Based on our clarification, I renamed this contributor in the documentation. The recorded repository work under the Efe Girgin identity includes user-facing content and feature work such as the FAQ page, refund-policy updates, and the Spotlight event page together with supporting styling and query updates."),
        ],
    )

    add_text_paragraph(doc, "Known Notes", style="Heading 1")
    add_bullets(
        doc,
        [
            "The project relies on an active MySQL instance and environment configuration before data-backed features can be demonstrated.",
            "Production builds may depend on network availability when external Google fonts are fetched during build-time.",
            "The documentation in this report is based on the current repository state and its included setup instructions.",
        ],
    )

    add_text_paragraph(doc, "Conclusion", style="Heading 1")
    add_text_paragraph(
        doc,
        "PulsePass demonstrates a complete multi-role event management workflow with secure authentication, structured server-side validation, transactional booking logic, and a clean component-based interface. The project balances technical implementation concerns with user-focused flows, making it a strong example of a full-stack software web development submission.",
        after=8,
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    create_document()
